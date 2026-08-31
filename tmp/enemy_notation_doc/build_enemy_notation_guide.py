from __future__ import annotations

import hashlib
import json
import shutil
import zipfile
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from urllib.parse import urlparse

from lxml import etree
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE


ROOT = Path(__file__).resolve().parents[2]
REFERENCE = ROOT / "docs" / "enemy-statistics-guide.docx"
FINAL = ROOT / "docs" / "enemy-defense-resistance-notation-guide.docx"
WORK = Path(__file__).resolve().parent
BODY_SOURCE = WORK / "body-source.docx"
INVENTORY = WORK / "package-diff.json"

EXPECTED_REFERENCE_SHA = "4D946B110F1A1DA51B7BF95EE845998A50D51C425FBBDD661A68D182FB8B0A8C"
ALLOWED_CHANGED_PARTS = {
    "word/document.xml",
    "word/_rels/document.xml.rels",
    "docProps/core.xml",
}

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CP_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
DC_NS = "http://purl.org/dc/elements/1.1/"
DCTERMS_NS = "http://purl.org/dc/terms/"
XSI_NS = "http://www.w3.org/2001/XMLSchema-instance"

INK = "222222"
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6570"
LIGHT = "F4F6F9"
PALE_BLUE = "E8EEF5"
BORDER = "D6DEE8"
WHITE = "FFFFFF"


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest().upper()


def sha256_file(path: Path) -> str:
    return sha256_bytes(path.read_bytes())


def set_run_font(run, *, latin="Calibri", east_asia="Yu Gothic", size=11,
                 color=INK, bold=False, italic=False):
    run.font.name = latin
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "ja-JP")
    lang.set(qn("w:eastAsia"), "ja-JP")
    return run


def set_paragraph_spacing(paragraph, *, before=0, after=6, line=1.25,
                          keep_with_next=False, keep_together=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_with_next
    fmt.keep_together = keep_together


def add_body_paragraph(doc, text="", *, after=6, before=0, line=1.25,
                       bold_prefix=None, color=INK, size=11, keep_together=False):
    p = doc.add_paragraph()
    set_paragraph_spacing(
        p, before=before, after=after, line=line, keep_together=keep_together
    )
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), size=size, color=color, bold=True)
        set_run_font(p.add_run(text[len(bold_prefix):]), size=size, color=color)
    else:
        set_run_font(p.add_run(text), size=size, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    if level == 1:
        set_paragraph_spacing(p, before=18, after=10, line=1.0, keep_with_next=True)
        set_run_font(p.add_run(text), size=16, color=BLUE, bold=True)
    else:
        set_paragraph_spacing(p, before=14, after=7, line=1.0, keep_with_next=True)
        set_run_font(p.add_run(text), size=13, color=BLUE, bold=True)
    return p


def set_numbering(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    for old in p_pr.findall(qn("w:numPr")):
        p_pr.remove(old)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)


def add_list_item(doc, text, *, numbered=False, after=4, size=11):
    p = doc.add_paragraph()
    set_numbering(p, 11 if numbered else 10)
    set_paragraph_spacing(p, before=0, after=after, line=1.25, keep_together=True)
    set_run_font(p.add_run(text), size=size)
    return p


def set_cell_margins(cell, top=100, bottom=100, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_border(cell, *, left=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    if left:
        edge = borders.find(qn("w:left"))
        if edge is None:
            edge = OxmlElement("w:left")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(left.get("sz", 18)))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), left.get("color", BLUE))


def set_table_borders(table, color=BORDER, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def set_table_geometry(table, widths, *, indent=120, borders=True):
    if sum(widths) != 9360:
        raise ValueError(f"Table widths must total 9360 DXA: {widths}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[i]))
            cell.width = Inches(widths[i] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if borders:
        set_table_borders(table)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    hdr = OxmlElement("w:tblHeader")
    hdr.set(qn("w:val"), "true")
    tr_pr.append(hdr)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def format_cell_paragraph(paragraph, *, align=WD_ALIGN_PARAGRAPH.LEFT, after=0, line=1.15):
    paragraph.alignment = align
    set_paragraph_spacing(paragraph, before=0, after=after, line=line, keep_together=True)


def add_callout(doc, label, body, *, formula_lines=None):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent=170, borders=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    set_cell_border(cell, left={"sz": 22, "color": BLUE})
    set_cell_margins(cell, top=130, bottom=130, start=170, end=150)
    p = cell.paragraphs[0]
    clear_paragraph(p)
    format_cell_paragraph(p, after=4, line=1.15)
    set_run_font(p.add_run(label), size=10, color=DARK_BLUE, bold=True)
    body_p = cell.add_paragraph()
    format_cell_paragraph(body_p, after=0, line=1.25)
    set_run_font(body_p.add_run(body), size=10.5, color=INK)
    if formula_lines:
        for line in formula_lines:
            fp = cell.add_paragraph()
            format_cell_paragraph(fp, after=2, line=1.1)
            set_run_font(fp.add_run(line), latin="Consolas", east_asia="Yu Gothic",
                         size=10.2, color=NAVY, bold=True)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=2, line=1.0)
    return table


def add_hyperlink(paragraph, text, url):
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Yu Gothic")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "19")
    run_props.extend([fonts, color, underline, size])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([run_props, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return rel_id


def add_source_item(doc, lead, link_text=None, url=None, tail=""):
    p = doc.add_paragraph()
    set_numbering(p, 10)
    set_paragraph_spacing(p, before=0, after=3, line=1.15, keep_together=True)
    if lead:
        set_run_font(p.add_run(lead), size=9.5, color=INK)
    if link_text and url:
        add_hyperlink(p, link_text, url)
    if tail:
        set_run_font(p.add_run(tail), size=9.5, color=INK)
    return p


def build_body_source():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, before, after in ((1, 16, 18, 10), (2, 13, 14, 7)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Guide Subtitle" not in [s.name for s in doc.styles]:
        doc.styles.add_style("Guide Subtitle", WD_STYLE_TYPE.PARAGRAPH)

    # Title block
    kicker = doc.add_paragraph()
    set_paragraph_spacing(kicker, before=0, after=5, line=1.0)
    set_run_font(kicker.add_run("ENEMY ANALYSIS / DISPLAY REFERENCE"),
                 size=9.5, color=BLUE, bold=True)

    title = doc.add_paragraph()
    set_paragraph_spacing(title, before=0, after=4, line=1.0, keep_with_next=True)
    set_run_font(title.add_run("敵の防御力・術耐性 表記ガイド"),
                 size=28, color=NAVY, bold=True)

    subtitle = doc.add_paragraph(style="Guide Subtitle")
    set_paragraph_spacing(subtitle, before=0, after=7, line=1.1, keep_with_next=True)
    set_run_font(subtitle.add_run("段階評価と実数値の対応・ダメージ計算での読み方"),
                 size=13.5, color=MUTED)

    meta = doc.add_paragraph()
    set_paragraph_spacing(meta, before=0, after=14, line=1.0)
    set_run_font(meta.add_run("対象画面: Enemy Analysis   |   Version 1.0   |   2026-08-31"),
                 size=9.5, color=MUTED)

    add_callout(
        doc,
        "最重要",
        "ゲーム内の文字ランクは、敵の強さをおおまかに示す範囲表記です。正確なダメージ比較には実数値を使い、防御力と術耐性は別の仕組みとして読みます。",
    )

    add_heading(doc, "1. 表示ランクの意味", 1)
    add_body_paragraph(
        doc,
        "改修後の敵図鑑では、防御力・術耐性を E、D、C、B、B+、A、A+、S、S+、SS の10段階で表示します。文字が上がるほど数値も高くなりますが、画面上のランクだけでは厳密な値は分かりません。",
    )
    add_list_item(doc, "Enemy Analysis の初期表示は「ゲーム内評価」です。「実数値」へ切り替えると、比較に用いる基礎値を確認できます。")
    add_list_item(doc, "本アプリは図鑑の評価文字列を直接読むのではなく、取得した実数値をゲーム内と同じ段階基準として換算します。")
    add_list_item(doc, "値が欠損・未定義・非数値の場合は評価できないため「—」と表示します。")

    add_callout(
        doc,
        "10段階の順序",
        "左ほど低く、右ほど高い評価です。「+」を含む文字は独立した段階として扱います。",
        formula_lines=["E < D < C < B < B+ < A < A+ < S < S+ < SS"],
    )

    pbreak = doc.add_paragraph()
    pbreak.add_run().add_break(WD_BREAK.PAGE)

    add_heading(doc, "2. 本アプリの換算条件", 1)
    mapping = [
        ("SS", "5000 < x", "整数: 5001以上", "90 < x", "整数: 91以上"),
        ("S+", "3000 ≤ x ≤ 5000", "整数: 3000～5000", "80 ≤ x ≤ 90", "整数: 80～90"),
        ("S", "2000 ≤ x < 3000", "整数: 2000～2999", "70 ≤ x < 80", "整数: 70～79"),
        ("A+", "1200 ≤ x < 2000", "整数: 1200～1999", "60 ≤ x < 70", "整数: 60～69"),
        ("A", "1000 ≤ x < 1200", "整数: 1000～1199", "50 ≤ x < 60", "整数: 50～59"),
        ("B+", "800 ≤ x < 1000", "整数: 800～999", "30 ≤ x < 50", "整数: 30～49"),
        ("B", "500 ≤ x < 800", "整数: 500～799", "20 ≤ x < 30", "整数: 20～29"),
        ("C", "200 ≤ x < 500", "整数: 200～499", "10 ≤ x < 20", "整数: 10～19"),
        ("D", "100 ≤ x < 200", "整数: 100～199", "0 < x < 10", "整数: 1～9"),
        ("E", "x < 100", "通常の整数値: 0～99", "x ≤ 0", "通常の値: 0"),
    ]
    table = doc.add_table(rows=1, cols=3)
    header = table.rows[0]
    for cell, text in zip(header.cells, ("評価", "防御力（DEF）", "術耐性（RES）")):
        set_cell_shading(cell, DARK_BLUE)
        p = cell.paragraphs[0]
        clear_paragraph(p)
        format_cell_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_run_font(p.add_run(text), size=9.5, color=WHITE, bold=True)
    set_repeat_table_header(header)

    for rating, def_cond, def_int, res_cond, res_int in mapping:
        cells = table.add_row().cells
        p = cells[0].paragraphs[0]
        clear_paragraph(p)
        format_cell_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_run_font(p.add_run(rating), size=10.5, color=NAVY, bold=True)
        for cell, cond, integer_text in ((cells[1], def_cond, def_int), (cells[2], res_cond, res_int)):
            p = cell.paragraphs[0]
            clear_paragraph(p)
            format_cell_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.05)
            set_run_font(p.add_run(cond), latin="Consolas", east_asia="Yu Gothic",
                         size=9.3, color=INK, bold=True)
            p.add_run().add_break()
            set_run_font(p.add_run(integer_text), size=8.4, color=MUTED)
    set_table_geometry(table, [900, 4230, 4230])

    note = add_body_paragraph(
        doc,
        "注: 公式にランク境界の数値表はないため、本表は2026-08-31時点の本アプリ実装値です。判定は丸め前の値に対して行い、整数の目安は読みやすさのために併記しています。境界上では、防御力5000・術耐性90はいずれもS+で、それを超えるとSSです。",
        before=4,
        after=0,
        line=1.15,
        size=8.8,
        color=MUTED,
        keep_together=True,
    )

    pbreak = doc.add_paragraph()
    pbreak.add_run().add_break(WD_BREAK.PAGE)

    # Page 2
    add_heading(doc, "3. 防御力と術耐性は、同じ文字でも働きが違う", 1)
    add_body_paragraph(
        doc,
        "防御力は物理ダメージから固定値を差し引き、術耐性は術ダメージを割合で減らします。そのため、同じB+やAでも軽減量が等しいとは限りません。",
        after=8,
    )

    comp = doc.add_table(rows=4, cols=2)
    for i, text in enumerate(("物理ダメージ / 防御力", "術ダメージ / 術耐性")):
        set_cell_shading(comp.rows[0].cells[i], PALE_BLUE)
        p = comp.rows[0].cells[i].paragraphs[0]
        clear_paragraph(p)
        format_cell_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_run_font(p.add_run(text), size=10, color=NAVY, bold=True)
    set_repeat_table_header(comp.rows[0])
    rows = [
        ("方式", "1ヒットごとの固定値減算", "割合による軽減"),
        ("基本式", "max(ATK − DEF, ATK × 5%)", "max(ATK × (1 − RES/100), ATK × 5%)"),
        ("読み方", "攻撃力と防御力の差が重要", "RES 1につき約1ポイント軽減率が上がる"),
    ]
    for row_i, (label, left, right) in enumerate(rows, start=1):
        for col_i, text in enumerate((left, right)):
            cell = comp.rows[row_i].cells[col_i]
            p = cell.paragraphs[0]
            clear_paragraph(p)
            format_cell_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.15)
            if row_i == 1:
                set_run_font(p.add_run(text), size=9.8, color=INK)
            elif row_i == 2:
                set_run_font(p.add_run(text), latin="Consolas", east_asia="Yu Gothic",
                             size=9.2, color=NAVY, bold=True)
            else:
                set_run_font(p.add_run(text), size=9.4, color=INK)
    set_table_geometry(comp, [4680, 4680])

    add_heading(doc, "同じB+の例", 2)
    add_callout(
        doc,
        "例: 攻撃力1000で攻撃",
        "敵の防御力800、術耐性30は、どちらもB+です。",
        formula_lines=[
            "物理: max(1000 − 800, 50) = 200",
            "術　: max(1000 × (1 − 0.30), 50) = 700",
            "結論: 同じランクでも、受けるダメージは一致しない",
        ],
    )

    add_list_item(doc, "物理攻撃は1ヒット攻撃力が低いほど防御力に止められやすく、高いほど防御力を上回った分を通しやすくなります。", after=3, size=10.3)
    add_list_item(doc, "術耐性は割合軽減なので、同じRESなら攻撃力が変わっても軽減率は原則同じです。", after=3, size=10.3)
    add_list_item(doc, "確定ダメージは、防御力・術耐性の影響を受けません。", after=3, size=10.3)

    add_heading(doc, "4. Enemy Analysis での使い分け", 1)
    add_list_item(doc, "「ゲーム内評価」で、物理と術の通りやすさを大まかに判断する。", numbered=True, after=2, size=10.2)
    add_list_item(doc, "「実数値」または敵詳細で、境界内のどこにいるかを確認する。", numbered=True, after=2, size=10.2)
    add_list_item(doc, "ダメージ計算機へ実数値を入力し、1ヒット・DPS・総ダメージを比べる。", numbered=True, after=2, size=10.2)
    add_list_item(doc, "ステージ補正、形態、自己強化、デバフや無視効果がある場合は基礎値だけで結論を出さない。", numbered=True, after=2, size=10.2)

    pbreak = doc.add_paragraph()
    pbreak.add_run().add_break(WD_BREAK.PAGE)

    # Page 3
    add_heading(doc, "5. 表示値の前提と例外", 1)
    definitions = [
        ("基礎値: ", "JP版ゲームデータの敵図鑑と戦闘データを敵IDで結合し、防御力 def と術耐性 magicResistance を読み取ります。複数レベルでは原則レベル0を採用します。"),
        ("ステージ差: ", "ゲーム内のSp.は、そのステージの敵が図鑑の標準情報と異なることを示します。本アプリはステージ固有条件やイベント補正を反映しません。"),
        ("形態・能力: ", "形態変化、自己強化、強襲条件、契約、デバフや防御・術耐性無視により、戦闘中の実効値は変わります。"),
        ("境界と丸め: ", "評価は丸め前の値で判定します。境界付近に小数値があると、整数表示との見た目が直感に合わない場合があります。"),
    ]
    for label, body in definitions:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=4, line=1.2, keep_together=True)
        set_run_font(p.add_run(label), size=10.3, color=NAVY, bold=True)
        set_run_font(p.add_run(body), size=10.3, color=INK)

    add_callout(
        doc,
        "用語の混同に注意",
        "術耐性（RES）は術ダメージを軽減する値です。スタン・睡眠などへの状態異常耐性、元素耐性、損傷耐性とは別の項目です。",
    )

    add_heading(doc, "6. 読み取りチェックリスト", 1)
    checks = [
        "文字ランクは正確な数値ではなく、範囲だと理解している。",
        "防御力は固定値減算、術耐性は割合軽減として別々に読んでいる。",
        "実数値、Sp.、形態・能力説明、ステージ固有補正を確認している。",
        "最終判断はダメージ計算機で実数値を使っている。",
    ]
    for check in checks:
        add_list_item(doc, check, numbered=False, after=2, size=10)

    add_heading(doc, "7. 参照情報", 1)
    add_body_paragraph(
        doc,
        "実装根拠: src/lib/enemyData.ts（換算・データ取得）、src/lib/damageCalculator.ts（基本式）、src/components/EnemyAnalysis.tsx（表示切替・基礎値注記）。",
        after=5,
        line=1.15,
        size=9.5,
        color=MUTED,
        keep_together=True,
    )
    add_source_item(
        doc,
        "公式: ",
        "アークナイツ公式 アプリ更新のお知らせ",
        "https://www.arknights.jp/news/377",
        "（2023-12-21の敵図鑑仕様変更）",
    )
    add_source_item(
        doc,
        "開発側告知: ",
        "敵方档案库优化内容前瞻",
        "https://www.taptap.cn/moment/415548477913697519",
        "（表示項目、Sp.、形態・耐性説明）",
    )
    add_source_item(
        doc,
        "補助資料: ",
        "Arknights Terra Wiki - DEF / RES / UI UX Changes",
        "https://arknights.wiki.gg/wiki/UI_UX_Changes",
        "（10段階表示と改修内容の整理）",
    )
    add_source_item(
        doc,
        "計算式: ",
        "Arknights Terra Wiki - Physical damage / Arts damage",
        "https://arknights.wiki.gg/wiki/Damage",
        "（減算・割合軽減・最低保証）",
    )
    doc.save(BODY_SOURCE)


def extract_body_and_hyperlinks():
    with zipfile.ZipFile(BODY_SOURCE, "r") as z:
        document_xml = z.read("word/document.xml")
        rels_xml = z.read("word/_rels/document.xml.rels")
    doc_root = etree.fromstring(document_xml)
    body = doc_root.find(f"{{{W_NS}}}body")
    if body is None:
        raise RuntimeError("Generated document has no body")
    body_children = [deepcopy(child) for child in list(body) if child.tag != f"{{{W_NS}}}sectPr"]

    rel_root = etree.fromstring(rels_xml)
    hyperlink_rels = {}
    for rel in rel_root:
        rel_type = rel.get("Type", "")
        if rel_type.endswith("/hyperlink") and rel.get("TargetMode") == "External":
            hyperlink_rels[rel.get("Id")] = rel.get("Target")
    return body_children, hyperlink_rels


def next_relationship_ids(ref_rel_root, count):
    used = set()
    for rel in ref_rel_root:
        rid = rel.get("Id", "")
        if rid.startswith("rId") and rid[3:].isdigit():
            used.add(int(rid[3:]))
    value = max(used, default=0) + 1
    result = []
    while len(result) < count:
        if value not in used:
            result.append(f"rId{value}")
        value += 1
    return result


def patch_document_xml(ref_xml: bytes, body_children, rid_map):
    root = etree.fromstring(ref_xml)
    body = root.find(f"{{{W_NS}}}body")
    if body is None:
        raise RuntimeError("Reference document has no body")
    sect_pr = body.find(f"{{{W_NS}}}sectPr")
    if sect_pr is None:
        raise RuntimeError("Reference document has no section properties")
    saved_sect_pr = deepcopy(sect_pr)
    for child in list(body):
        body.remove(child)
    for child in body_children:
        for hyperlink in child.iter(f"{{{W_NS}}}hyperlink"):
            old_id = hyperlink.get(f"{{{R_NS}}}id")
            if old_id in rid_map:
                hyperlink.set(f"{{{R_NS}}}id", rid_map[old_id])
        body.append(child)
    body.append(saved_sect_pr)
    return etree.tostring(
        root,
        encoding="UTF-8",
        xml_declaration=True,
        standalone=True,
    )


def patch_relationships(ref_xml: bytes, hyperlink_rels):
    root = etree.fromstring(ref_xml)
    new_ids = next_relationship_ids(root, len(hyperlink_rels))
    rid_map = {}
    for (old_id, target), new_id in zip(hyperlink_rels.items(), new_ids):
        if urlparse(target).scheme not in {"http", "https"}:
            raise ValueError(f"Unexpected hyperlink target: {target}")
        rel = etree.Element(f"{{{PKG_REL_NS}}}Relationship")
        rel.set("Id", new_id)
        rel.set(
            "Type",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        )
        rel.set("Target", target)
        rel.set("TargetMode", "External")
        root.append(rel)
        rid_map[old_id] = new_id
    return etree.tostring(
        root,
        encoding="UTF-8",
        xml_declaration=True,
        standalone=True,
    ), rid_map


def patch_core_properties(ref_xml: bytes):
    root = etree.fromstring(ref_xml)
    title = root.find(f"{{{DC_NS}}}title")
    if title is None:
        title = etree.SubElement(root, f"{{{DC_NS}}}title")
    title.text = "敵の防御力・術耐性 ゲーム内表記ガイド"
    subject = root.find(f"{{{DC_NS}}}subject")
    if subject is None:
        subject = etree.SubElement(root, f"{{{DC_NS}}}subject")
    subject.text = "Enemy Analysis の段階評価と実数値の対応"
    modified = root.find(f"{{{DCTERMS_NS}}}modified")
    if modified is not None:
        modified.text = datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")
        modified.set(f"{{{XSI_NS}}}type", "dcterms:W3CDTF")
    return etree.tostring(
        root,
        encoding="UTF-8",
        xml_declaration=True,
        standalone=True,
    )


def build_final():
    if sha256_file(REFERENCE) != EXPECTED_REFERENCE_SHA:
        raise RuntimeError("Reference DOCX hash does not match artifact.md")
    build_body_source()
    body_children, hyperlink_rels = extract_body_and_hyperlinks()

    with zipfile.ZipFile(REFERENCE, "r") as src:
        ref_rels = src.read("word/_rels/document.xml.rels")
        new_rels, rid_map = patch_relationships(ref_rels, hyperlink_rels)
        replacements = {
            "word/document.xml": patch_document_xml(
                src.read("word/document.xml"), body_children, rid_map
            ),
            "word/_rels/document.xml.rels": new_rels,
            "docProps/core.xml": patch_core_properties(src.read("docProps/core.xml")),
        }
        with zipfile.ZipFile(FINAL, "w") as dst:
            for item in src.infolist():
                data = replacements.get(item.filename, src.read(item.filename))
                dst.writestr(item, data)

    if sha256_file(REFERENCE) != EXPECTED_REFERENCE_SHA:
        raise RuntimeError("Reference DOCX changed during authoring")

    with zipfile.ZipFile(REFERENCE, "r") as ref_zip, zipfile.ZipFile(FINAL, "r") as out_zip:
        ref_parts = {name: sha256_bytes(ref_zip.read(name)) for name in ref_zip.namelist()}
        out_parts = {name: sha256_bytes(out_zip.read(name)) for name in out_zip.namelist()}
    changed = sorted(
        name for name in set(ref_parts) | set(out_parts)
        if ref_parts.get(name) != out_parts.get(name)
    )
    unexpected = sorted(set(changed) - ALLOWED_CHANGED_PARTS)
    report = {
        "reference": str(REFERENCE),
        "reference_sha256": sha256_file(REFERENCE),
        "final": str(FINAL),
        "final_sha256": sha256_file(FINAL),
        "changed_parts": changed,
        "unexpected_changed_parts": unexpected,
        "reference_part_count": len(ref_parts),
        "final_part_count": len(out_parts),
    }
    INVENTORY.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    if unexpected:
        raise RuntimeError(f"Unexpected package changes: {unexpected}")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    build_final()
