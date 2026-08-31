from docx import Document

document = Document()
document.add_heading('Render smoke test', level=1)
document.add_paragraph('PDF export check.')
document.save(r'C:\Users\legen\Documents\arknights_2\tmp\enemy_notation_doc\smoke.docx')
